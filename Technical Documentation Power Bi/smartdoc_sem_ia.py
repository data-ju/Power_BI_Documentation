"""
Automação para Geração de Documentação de Projetos Power BI
-----------------------------------------------------------

Este código automatiza a documentação de relatórios Power BI a partir de um arquivo `.pbit` 
convertido em `.zip`. Ele extrai informações de páginas, tabelas, colunas, medidas, fontes, 
e relacionamentos para gerar uma documentação detalhada em Word.

Configurações necessárias:
1. Converter o arquivo `.pbix` em `.pbit` (via Power BI Desktop).
2. Configurar caminhos e nomes dos arquivos no módulo `config`.
3. Instalar dependências listadas em `requirements.txt`.

"""

"""
Primeira parte do código
1. Leitura e obtenção das informações a partir dos arquivos "Layout" e "ModelSchema" em JSON
"""

import json
import config as cfg
import os 
from os import path, rename
import zipfile
from datetime import datetime
import requests
from docx import Document
from io import StringIO
import time

# #Instalando as bibliotecas
# !pip install -r requirements.txt

# Função para verificar e renomear arquivos
def verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip):
    if os.path.exists(arquivo_zip):
        print("Arquivo .zip já existe. Pulando para a próxima instrução.")
    else:
        os.rename(arquivo_pbit, arquivo_zip)

# Função para extrair arquivos do ZIP
def extrair_arquivos_zip(arquivo_zip, caminho_BI, arquivos_para_extrair):
    with zipfile.ZipFile(arquivo_zip, 'r') as zip_ref:
        for arquivo in arquivos_para_extrair:
            zip_ref.extract(arquivo, caminho_BI)

# Função para carregar os dados do JSON
def carregar_dados_json(arquivo: str, encoding: str = 'utf-16-le') -> dict:
    """Carrega dados de um arquivo JSON."""
    try:
        with open(arquivo, 'r', encoding=encoding) as f:
            return json.load(f)
    except Exception as e:
        print(f"Erro ao carregar JSON: {arquivo} - {e}")
        return {}

# Função para extrair as páginas do arquivo "Layout"
def extrair_paginas(layout: dict) -> str:
    """Extrai e organiza informações de páginas em formato Markdown."""
    markdown_output = [""]
    for section in layout.get('sections', []):
        page_name = section.get('displayName', 'Sem Nome')
        markdown_output.append(f"{page_name}\n-----------\n")
    return "\n".join(markdown_output)

# Função para extrair os visuais do arquivo "Layout"
def extrair_visuais(layout: dict) -> str:
    """Extrai e organiza informações de visuais em cada página em formato Markdown."""
    markdown_output = [""]
    for section in layout.get('sections', []):
        page_name = section.get('displayName', 'Sem Nome')
        for container in section.get("visualContainers", []):
            config_data = json.loads(container.get("config", "{}"))
            visual_type = config_data.get("singleVisual", {}).get("visualType")
            position = next(iter(config_data.get("layouts", [])), {}).get("position", {})
            query_refs = [item.get("queryRef") for items in config_data.get("singleVisual", {}).get("projections", {}).values()
                          for item in items if item.get("queryRef")]
            
            markdown_output.append(
                f"Página: {page_name}\n"
                f"X: {int(position.get('x', 0))}\n"
                f"Y: {int(position.get('y', 0))}\n"
                f"Altura: {int(position.get('height', 0))}\n"
                f"Largura: {int(position.get('width', 0))}\n"
                f"Tipo de visual: {visual_type}\n"
                f"Medidas utilizadas: {', '.join(query_refs) if query_refs else 'Não há medidas utilizadas no visual'}\n"
                "-----------\n"
            )
    return "\n".join(markdown_output)

def extrair_tabelas(model_data: dict) -> str:
    """Extrai e organiza informações de tabelas em formato Markdown."""
    markdown_output = [""]
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        if table_name.startswith("DateTableTemplate") or table_name.startswith("LocalDateTable"):
            continue
        for column in table.get('columns', []):
            column_name = column.get("name", "")
            data_type = column.get('dataType', "")
            is_calculated = 'Sim' if column.get('type', "") in ['calculatedTableColumn', 'calculated'] else 'Não'
            markdown_output.append(
                f"Tabela: {table_name}\n"
                f"Coluna: {column_name}\n"
                f"Tipo de dados: {data_type}\n"
                f"Coluna calculada?: {is_calculated}\n"
                "-----------\n"
            )
    return "\n".join(markdown_output)

def extrair_medidas(model_data: dict) -> str:
    """Extrai e organiza informações de medidas em formato Markdown."""
    markdown_output = [""]
    processed_measures = set()
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        for measure in table.get('measures', []):
            measure_name = measure.get('name', '')
            measure_expression = measure.get('expression', '')
            if (table_name, measure_name) in processed_measures:
                continue
            processed_measures.add((table_name, measure_name))
            if isinstance(measure_expression, list):
                measure_expression = ' '.join(filter(lambda x: x.strip(), measure_expression))
            markdown_output.append(
                f"Tabela: {table_name}\n"
                f"Medida: {measure_name}\n"
                f"Expressão: {measure_expression}\n"
                "-----------\n"
            )
    return "\n".join(markdown_output)

def extrair_fontes(model_data: dict) -> str:
    """Extrai e organiza informações sobre fontes de dados em formato Markdown."""
    markdown_output = [""]
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        if table_name.startswith("DateTableTemplate") or table_name.startswith("LocalDateTable"):
            continue
        for partition in table.get('partitions', []):
            partition_mode = partition.get('mode')
            source = partition.get('source', {})
            font_type = source.get('type')
            font_expression = source.get('expression')
            if isinstance(font_expression, list):
                font_expression = ' '.join(filter(lambda x: x.strip(), font_expression))
            markdown_output.append(
                f"Tabela: {table_name}\n"
                f"Modo de importação: {partition_mode}\n"
                f"Tipo de importação: {font_type}\n"
                f"Fonte: {font_expression}\n"
                "-----------\n"
            )
    return "\n".join(markdown_output)

def extrair_relacionamentos(model_data: dict) -> str:
    """Extrai e organiza informações de relacionamentos em formato Markdown."""
    markdown_output = [""]
    for relation in model_data.get('model', {}).get('relationships', []):
        from_table = relation.get('fromTable')
        to_table = relation.get('toTable')
        from_column = relation.get('fromColumn', '')
        to_column = relation.get('toColumn', '')
        if from_table.startswith("DateTableTemplate") or from_table.startswith("LocalDateTable") or \
           to_table.startswith("DateTableTemplate") or to_table.startswith("LocalDateTable"):
            continue
        markdown_output.append(
            f"Da tabela: {from_table}\n"
            f"Para tabela: {to_table}\n"
            f"Da coluna: {from_column}\n"
            f"Para coluna: {to_column}\n"
            "-----------\n"
        )
    return "\n".join(markdown_output)

"""
Segunda parte do código: 
1. Exportação para o Word das informações em Markdown
"""

def salvar_versao(salvar_path):
    """
    Verifica se o arquivo já existe e, se sim, cria um nome de arquivo com uma versão incremental.
    
    Parâmetros:
        salvar_path (str): Caminho completo do arquivo que deseja salvar.
        
    Retorna:
        str: Caminho final do arquivo com uma versão incrementada, se necessário.
    """

    # Se o arquivo não existir, retorna o caminho original
    if not os.path.exists(salvar_path):
        return salvar_path
    
    # Se já existir, adiciona a versão incremental
    base, ext = os.path.splitext(salvar_path)
    versao = 2
    
    # Incrementa a versão até encontrar um nome de arquivo disponível
    while os.path.exists(f"{base}_versão_{versao:02}{ext}"):
        versao += 1
    
    # Retorna o novo caminho do arquivo com a versão adicionada
    return f"{base}_versão_{versao:02}{ext}"

# Função para gerar o documento com conteúdo Markdown
def gerar_documento(cfg, extracoes, modelo_path, salvar_path):
    """
    Gera o documento Word com as descrições em formato Markdown nos locais apropriados.
    """
    # Carrega o modelo de documento do Word
    document = Document(modelo_path)

    # Preenche informações básicas do documento
    for para in document.paragraphs:
        if "Data da documentação:" in para.text:
            para.add_run(f" {datetime.now().strftime('%d/%m/%Y')}")
        elif "Nome do Relatório:" in para.text:
            para.add_run(f" {cfg.nome_BI}")

    # Insere cada seção Markdown no local correto do documento
    for titulo, conteudo_markdown in extracoes.items():
        for para in document.paragraphs:
            if para.text.strip() == titulo.capitalize():
                # Insere o conteúdo Markdown logo abaixo do parágrafo do título
                paragrafo_conteudo = document.add_paragraph(conteudo_markdown)
                para._element.addnext(paragrafo_conteudo._element)
                break  

    # Gera e salva o documento no caminho final, com controle de versão se necessário
    caminho_final = salvar_versao(salvar_path)
    document.save(caminho_final)
    print(f'Documentação gerada com sucesso em: {caminho_final}')

# Função principal para execução do processo
def main():
    caminho_BI, nome_BI = cfg.caminho_BI, cfg.nome_BI
    modelo_path = os.path.join(cfg.caminho_modelo_word, cfg.nome_modelo_word)
    salvar_path = os.path.join(cfg.caminho_documentação, f'{nome_BI}_doc.docx')
    arquivo_pbit = os.path.join(caminho_BI, f'{nome_BI}.pbit')
    arquivo_zip = os.path.join(caminho_BI, f'{nome_BI}.zip')

    # Verifica e renomeia o arquivo para .zip, se necessário
    verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip)

    # Extrai arquivos do ZIP
    extrair_arquivos_zip(arquivo_zip, caminho_BI, ['Report/Layout', 'DataModelSchema'])

    # Carrega os dados JSON
    layout_data = carregar_dados_json(os.path.join(caminho_BI, 'Report/Layout'))
    model_data = carregar_dados_json(os.path.join(caminho_BI, 'DataModelSchema'))

    # Reverte o arquivo ZIP para o formato original .pbit
    os.rename(arquivo_zip, arquivo_pbit)

    # Dicionário de extrações em Markdown
    extracoes = {
        "Páginas": extrair_paginas(layout_data),
        "Tabelas": extrair_tabelas(model_data),
        "Medidas": extrair_medidas(model_data),
        "Visuais": extrair_visuais(layout_data),
        "Fontes": extrair_fontes(model_data),
        "Relacionamentos": extrair_relacionamentos(model_data)
    }

    gerar_documento(cfg, extracoes, modelo_path, salvar_path)

if __name__ == '__main__':
    main()